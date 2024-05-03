from docx import Document
import easygui as gui

# Let the use choose the file
file = gui.fileopenbox()

# Import the document
doc = Document(file)

#Creating the new document
new_doc = Document()

# Initialize variables
talker = ""
text = ""


# Iterate through paragraphs
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text == "":
        break


    if i % 2 == 0:
        talker = paragraph.text.split(" ")[1]
    else:
        new_doc.add_paragraph(talker + ": " + paragraph.text)



# Save the changes
new_doc.save(file.split(".")[:-1][0] + "_edited.docx")
